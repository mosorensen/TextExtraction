# -*- coding: utf-8 -*-
"""
This script reads in one word document.
It then reads in all of the tables in that word document.
It then creates a dictionary for each table, and and overall dictionary
    connecting the doc_id to these document specific dictionaries per table
It also creates a dataframe with a row per document and with entries
    corresponding to the identified variables in all of the tables

To do:
    - Automatically detect which table is being interpreted. 
        - This can only be done once we look at many documents, as the tables 
            might vary by document. This could be implemented by checking if 
            the table contains certin key words, or by trying to find the 
            header immediately prior to the table.
                - Front Page Info
                - ghSMART Competency Ratings 
        - Only add variables and table dictionary for these tables
        - create seperate table parsers for each type of table being saved
    - Output the dataframe to an excel file. Preferably similar in structure
        to the excel file already created.
    - Try to detect other relevant information
        - Table of Contents -- to get an overview of frequency of various 
                data in the overall dataset
        - Key Strengths and Key Risks
        - Firm-Specific Mission Text
        - Firm-Specific Key Outcomes Ratings/Comments
        - Company Specific Questions
        - Candidate Specific Recommendations
        - Career Goals and Motivations
    - Transform the entire document into one parseable text. Then check
        to see how many times certain words occur in the text using
        dictionaries or arrays of strings. This can be a crude linguistic
        approach i.e. measuring the frequency of words associated with
        overconfidence (overconfident, assertive, optimistic, etc.) versus
        the frequency of words associated with caution (cautious, timid,
        scientific, etc.).

Notes:
    - Merged Cells are read as multiple cells all with the same information
    - info on the docx package: https://python-docx.readthedocs.io/en/latest/index.html
    - info on the pandas package: https://pandas.pydata.org/
"""


from docx import Document
import pandas as pd


comp_cat = ["intellectual ability", "personal effectiveness", 
            "team leadership", "interpersonal effectiveness", 
            "contextual skills/experience"]

document = Document("C:/Users/ydecress/Desktop/Cleaning_Small.docx")

tabs = document.tables
comp_tab = tabs[5]

for para in document.paragraphs:
    print(para.text)
    
# loops through the table and decides what table it is based on if any of
# its cells contains certain strings. To add robustness, we could
# check if multiple cells fulfill certain criteria.
def which_table(table):
    table_name="other"
    num_comp = 0
    for row in table.rows:
        last_text = ""
        for cell in row.cells:
            curr_text = cell.text
            #Check 1: does the table give identifying information
            if "Candidate Name" == curr_text.strip():
                table_name="info"
            #Check 2: does the table give competency information
            if curr_text.lower() in comp_cat:
                if last_text=="":
                    last_text = curr_text
                else:
                    if last_text==curr_text: #a merged cell is interpreted as two cells with the same contents
                        num_comp+=1 #keep track of the number of listed competencies
                    else:
                        last_text=curr_text #if not a merged cell, update last_text
    if num_comp == 5: #the table must identify all five competency categories
        table_name = "competencies"
            
    return table_name

i=which_table(comp_tab)

#To add robustness, we can add a variable to let the table parser know what type
#    of table it is currently dealing with.
def parse_table(table, doc_id):
    table_dict = {}
    table_data = pd.DataFrame({'doc_id':[doc_id]})
    for i, row in enumerate(table.rows): #loop through all of the rows
        text = (cell.text for cell in row.cells) #converts all tables cells to text  
        text = tuple(text) #need to convert from row object to a tuple
        j=0 #let j index over the entries in the row tuple
        while j < len(text)-1: #while we have at least one additional entry in the row
            key = text[j]
            if len(key) == 0: #if the current cell is empty, it can't be a key, so skip it
                j+=1
                continue
            entries = text[j+1] #the keys entry is next to it in the table
            j=j+2 #the next key is after the current key's entry...so two indices away
            if key==entries: #if the cell if a merged cell, skip it
                continue  
            # add the keys and entries into a table specific dictionary and dataframe
            table_dict[key] = entries 
            table_data[key] = entries     
    return (table_dict, table_data)    


tables_to_parse = ["info", "competencies"]
doc_paths = ["C:/Users/ydecress/Desktop/Cleaning_Small.docx", "C:/Users/ydecress/Desktop/Cleaning_Small_Copy.docx"]
for doc_id in range(len(doc_paths)):
    #first read in the given document
    print(doc_id)
    path = doc_paths[doc_id]
    document = Document(path)
    doc_dict = {} #dictionary of all tables
    doc_data = pd.DataFrame({'doc_id':[doc_id], 'doc_path':[path]})
    table_id = 0 #tracks the table index
    for table in document.tables:
        table_id += 1
        table_type = which_table(table)
        if table_type not in tables_to_parse:
            print(table_type, "not parsed")
            continue
        print(table_type, "parsing")
        table_dict, table_data = parse_table(table, doc_id) #parse the table and its dictionary and dataframe form
        # add the table to the document dictionary and document table
        doc_dict[table_type]=table_dict 
        doc_data = pd.merge(doc_data, table_data, on='doc_id', how='outer')
    candidate_name=doc_dict['info']['Candidate Name']
    #if this was the first document, we need to create the master dataframe
    #each row will be one executive and each column will contain a variable
    #each column gets a unique doc_id and path to document
    if doc_id==0: 
        data_master = doc_data
        dict_master = {candidate_name:doc_dict}
    else:
        data_master = data_master.append(doc_data)
        dict_master[candidate_name] = doc_dict


dict_master
data_master











